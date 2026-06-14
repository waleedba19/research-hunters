"""
verify_refs/reports.py — Build Excel (openpyxl) and DOCX (python-docx) reports.

Excel: master_database_verifyrefs.xlsx
  - Sheet "Summary":  totals + counts
  - Sheet "All References": one row per ref, color-coded by status
  - Sheet "Verified Details": full metadata of verified refs
  - Frozen header row, conditional formatting, professional column widths

DOCX: literature_verification_report.docx
  - Title page with date
  - Executive summary
  - Sections: Verified, Likely, Unverified, Fake (each with styled tables)
  - Page numbers, header/footer
  - Red highlight on UNVERIFIED / FAKE rows
"""
import os
from datetime import datetime
from typing import Dict, List, Any
from logger import get_logger

log = get_logger("verify_refs.reports")

# Status color palette
STATUS_COLORS = {
    "VERIFIED":   "C6EFCE",  # green
    "LIKELY":     "FFEB9C",  # yellow
    "UNVERIFIED": "FFC7CE",  # red
    "FAKE":       "C00000",  # dark red
}
STATUS_TEXT_COLORS = {
    "VERIFIED":   "006100",
    "LIKELY":     "9C5700",
    "UNVERIFIED": "9C0006",
    "FAKE":       "FFFFFF",
}


# ============================================================================
# EXCEL REPORT
# ============================================================================

def build_excel_report(classified: List[Dict[str, Any]], output_path: str,
                       source_description: str = "") -> str:
    """Build the Excel master database. Returns the output path."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        from openpyxl.formatting.rule import CellIsRule, FormulaRule
    except ImportError as e:
        log.error(f"openpyxl not installed: {e}")
        return ""

    wb = Workbook()

    # ---- Summary sheet ----
    ws = wb.active
    ws.title = "Summary"

    title_font = Font(name="Calibri", size=18, bold=True, color="FFFFFF")
    title_fill = PatternFill("solid", fgColor="1F4E78")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="2E75B6")
    thin = Side(border_style="thin", color="999999")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_wrap = Alignment(horizontal="left", vertical="center", wrap_text=True)

    # Title
    ws.merge_cells("A1:D1")
    ws["A1"] = "Literature Reference Verification Report"
    ws["A1"].font = title_font
    ws["A1"].fill = title_fill
    ws["A1"].alignment = center
    ws.row_dimensions[1].height = 32

    ws.merge_cells("A2:D2")
    ws["A2"] = f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
    ws["A2"].font = Font(name="Calibri", size=10, italic=True, color="555555")
    ws["A2"].alignment = center

    ws.merge_cells("A3:D3")
    ws["A3"] = f"Source: {source_description[:200]}"
    ws["A3"].font = Font(name="Calibri", size=10, italic=True, color="555555")
    ws["A3"].alignment = left_wrap

    # Counts table
    total = len(classified)
    counts = {"VERIFIED": 0, "LIKELY": 0, "UNVERIFIED": 0, "FAKE": 0}
    for c in classified:
        st = c.get("status", "FAKE")
        counts[st] = counts.get(st, 0) + 1
    pdfs = sum(1 for c in classified if c.get("download_success"))

    ws["A5"] = "Status"
    ws["B5"] = "Count"
    ws["C5"] = "Percentage"
    ws["D5"] = "Color"
    for cell in ws[5]:
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border

    row = 6
    for status, color in STATUS_COLORS.items():
        ws[f"A{row}"] = status
        ws[f"B{row}"] = counts.get(status, 0)
        ws[f"C{row}"] = f"{(counts.get(status, 0) / max(total, 1) * 100):.1f}%"
        ws[f"D{row}"] = color
        ws[f"A{row}"].font = Font(name="Calibri", size=11, bold=True, color=STATUS_TEXT_COLORS[status])
        ws[f"A{row}"].fill = PatternFill("solid", fgColor=color)
        ws[f"B{row}"].alignment = center
        ws[f"C{row}"].alignment = center
        for cell in ws[row]:
            cell.border = border
        row += 1

    ws[f"A{row}"] = "TOTAL"
    ws[f"B{row}"] = total
    ws[f"C{row}"] = "100.0%"
    for cell in ws[row]:
        cell.font = Font(name="Calibri", size=11, bold=True)
        cell.fill = PatternFill("solid", fgColor="D9D9D9")
        cell.border = border
        cell.alignment = center
    row += 1
    ws[f"A{row}"] = "PDFs Downloaded"
    ws[f"B{row}"] = pdfs
    for cell in ws[row]:
        cell.font = Font(name="Calibri", size=11, bold=True, color="1F4E78")
        cell.border = border
        cell.alignment = center

    # Column widths
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 12
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 12

    # ---- All References sheet ----
    ws2 = wb.create_sheet("All References")
    headers = [
        "#", "Status", "Score", "Reference (input)", "Matched Title",
        "Authors", "Year", "DOI", "Source", "Sources #", "PDF", "Notes / Error"
    ]
    for c, h in enumerate(headers, start=1):
        cell = ws2.cell(row=1, column=c, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border
    ws2.row_dimensions[1].height = 24
    ws2.freeze_panes = "A2"

    for i, c in enumerate(classified, start=1):
        row = i + 1
        status = c.get("status", "FAKE")
        color = STATUS_COLORS.get(status, "FFFFFF")
        text_color = STATUS_TEXT_COLORS.get(status, "000000")

        values = [
            i,
            status,
            round(float(c.get("score", 0.0)), 3),
            c.get("ref", "")[:500],
            c.get("matched_title", "")[:300],
            c.get("matched_authors", "")[:200],
            c.get("matched_year", ""),
            c.get("matched_doi", ""),
            c.get("source_platform", ""),
            c.get("source_count", 0),
            "YES" if c.get("download_success") else ("-" if c.get("status") != "VERIFIED" else "NO"),
            (c.get("reason", "") + " | " + c.get("error", "")).strip(" |")[:300],
        ]
        for col_idx, val in enumerate(values, start=1):
            cell = ws2.cell(row=row, column=col_idx, value=val)
            cell.fill = PatternFill("solid", fgColor=color)
            cell.font = Font(name="Calibri", size=10, color=text_color)
            cell.alignment = left_wrap
            cell.border = border
        ws2.row_dimensions[row].height = max(30, min(120, len(c.get("ref", "")) // 4))

    # Column widths
    widths = [4, 12, 7, 60, 50, 30, 8, 30, 18, 10, 6, 50]
    for c, w in enumerate(widths, start=1):
        ws2.column_dimensions[get_column_letter(c)].width = w

    # Auto filter
    ws2.auto_filter.ref = f"A1:L{len(classified) + 1}"

    # ---- Verified Details sheet ----
    ws3 = wb.create_sheet("Verified Details")
    headers3 = [
        "#", "Title", "Authors", "Year", "DOI", "Source Platform",
        "Sources #", "Cross-Source", "Score", "Reason", "PDF Downloaded", "PDF Path"
    ]
    for c, h in enumerate(headers3, start=1):
        cell = ws3.cell(row=1, column=c, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border
    ws3.row_dimensions[1].height = 24
    ws3.freeze_panes = "A2"

    verified = [c for c in classified if c.get("status") == "VERIFIED"]
    for i, c in enumerate(verified, start=1):
        row = i + 1
        values = [
            i,
            c.get("matched_title", "")[:300],
            c.get("matched_authors", "")[:200],
            c.get("matched_year", ""),
            c.get("matched_doi", ""),
            c.get("source_platform", ""),
            c.get("source_count", 0),
            "YES" if c.get("source_count", 0) >= 2 else "NO",
            round(float(c.get("score", 0.0)), 3),
            c.get("reason", "")[:300],
            "YES" if c.get("download_success") else "NO",
            c.get("download_path", ""),
        ]
        for col_idx, val in enumerate(values, start=1):
            cell = ws3.cell(row=row, column=col_idx, value=val)
            cell.fill = PatternFill("solid", fgColor="C6EFCE")
            cell.font = Font(name="Calibri", size=10, color="006100")
            cell.alignment = left_wrap
            cell.border = border
        ws3.row_dimensions[row].height = 30
    widths3 = [4, 50, 30, 8, 30, 18, 10, 12, 8, 50, 14, 50]
    for c, w in enumerate(widths3, start=1):
        ws3.column_dimensions[get_column_letter(c)].width = w

    # Save
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    wb.save(output_path)
    log.info(f"Excel report saved: {output_path}")
    return output_path


# ============================================================================
# DOCX REPORT
# ============================================================================

def build_docx_report(classified: List[Dict[str, Any]], output_path: str,
                      source_description: str = "") -> str:
    """Build a professionally-styled DOCX verification report."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
        from docx.enum.table import WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        log.error(f"python-docx not installed: {e}")
        return ""

    doc = Document()

    # ---- Page setup ----
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ---- Default style ----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Helper: cell shading ----
    def _shade_cell(cell, hex_color: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    def _add_badge(paragraph, text: str, hex_color: str, text_color: str = "FFFFFF") -> None:
        """Add a colored run that looks like a badge."""
        run = paragraph.add_run(f"  {text}  ")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(text_color)
        rPr = run._r.get_or_add_rPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        rPr.append(shd)

    # ---- Title page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("LITERATURE REFERENCE\nVERIFICATION REPORT")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string("1F4E78")
    title.paragraph_format.space_after = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"Generated on {datetime.utcnow().strftime('%B %d, %Y at %H:%M UTC')}")
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor.from_string("555555")
    sub.paragraph_format.space_after = Pt(12)

    if source_description:
        src = doc.add_paragraph()
        src.alignment = WD_ALIGN_PARAGRAPH.CENTER
        src_run = src.add_run(f"Source: {source_description[:200]}")
        src_run.font.size = Pt(10)
        src_run.font.color.rgb = RGBColor.from_string("888888")

    # ---- Executive Summary ----
    doc.add_page_break()
    h1 = doc.add_heading("Executive Summary", level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor.from_string("1F4E78")

    total = len(classified)
    counts = {"VERIFIED": 0, "LIKELY": 0, "UNVERIFIED": 0, "FAKE": 0}
    for c in classified:
        counts[c.get("status", "FAKE")] = counts.get(c.get("status", "FAKE"), 0) + 1
    pdfs = sum(1 for c in classified if c.get("download_success"))

    p = doc.add_paragraph()
    p.add_run(
        f"This report presents the verification results for {total} reference(s) extracted from the provided source. "
        f"Each reference was searched across 81 academic databases, scored for relevance using a local AI model, "
        f"and classified into one of four categories: "
    )
    p.add_run("VERIFIED ").bold = True
    p.add_run("(high-confidence match), ")
    p.add_run("LIKELY ").bold = True
    p.add_run("(moderate-confidence match), ")
    p.add_run("UNVERIFIED ").bold = True
    p.add_run("(candidates found but no strong match), and ")
    p.add_run("FAKE ").bold = True
    p.add_run("(no candidates found — likely fabricated or hallucinated).")

    # Stats table
    doc.add_paragraph()
    stats_p = doc.add_paragraph()
    stats_p.add_run("Verification Statistics").bold = True
    stats_p.runs[0].font.size = Pt(13)
    stats_p.runs[0].font.color.rgb = RGBColor.from_string("1F4E78")

    table = doc.add_table(rows=6, cols=3)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = "Status"
    hdr[1].text = "Count"
    hdr[2].text = "Percentage"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor.from_string("FFFFFF")
        _shade_cell(c, "1F4E78")

    row_data = [
        ("VERIFIED", counts["VERIFIED"], STATUS_COLORS["VERIFIED"], STATUS_TEXT_COLORS["VERIFIED"]),
        ("LIKELY", counts["LIKELY"], STATUS_COLORS["LIKELY"], STATUS_TEXT_COLORS["LIKELY"]),
        ("UNVERIFIED", counts["UNVERIFIED"], STATUS_COLORS["UNVERIFIED"], STATUS_TEXT_COLORS["UNVERIFIED"]),
        ("FAKE", counts["FAKE"], STATUS_COLORS["FAKE"], STATUS_TEXT_COLORS["FAKE"]),
        ("TOTAL", total, "D9D9D9", "000000"),
    ]
    for i, (status, n, bg, fg) in enumerate(row_data, start=1):
        cells = table.rows[i].cells
        cells[0].text = status
        cells[1].text = str(n)
        cells[2].text = f"{(n / max(total, 1) * 100):.1f}%"
        for c in cells:
            _shade_cell(c, bg)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.color.rgb = RGBColor.from_string(fg)
                    if i == 5 or status == "TOTAL":
                        r.font.bold = True

    # PDFs downloaded line
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run(f"PDFs successfully downloaded: ").bold = True
    p.add_run(f"{pdfs} of {counts['VERIFIED']} verified references")

    # ---- Section: Verified References ----
    if counts["VERIFIED"]:
        doc.add_page_break()
        h = doc.add_heading(f"1. Verified References ({counts['VERIFIED']})", level=1)
        for r in h.runs:
            r.font.color.rgb = RGBColor.from_string("1F4E78")
        doc.add_paragraph(
            "These references were matched with high confidence (≥85% AI score) and "
            "downloaded as PDFs where possible. They are recommended for inclusion "
            "in the literature review.")

        for i, c in enumerate([c for c in classified if c.get("status") == "VERIFIED"], 1):
            _add_reference_block(doc, c, i, "VERIFIED", _add_badge)

    # ---- Section: Likely References ----
    if counts["LIKELY"]:
        doc.add_page_break()
        h = doc.add_heading(f"2. Likely References ({counts['LIKELY']})", level=1)
        for r in h.runs:
            r.font.color.rgb = RGBColor.from_string("1F4E78")
        doc.add_paragraph(
            "These references had moderate AI score (60-85%). They may correspond to real "
            "publications but require manual verification. Review titles and DOIs to confirm.")
        for i, c in enumerate([c for c in classified if c.get("status") == "LIKELY"], 1):
            _add_reference_block(doc, c, i, "LIKELY", _add_badge)

    # ---- Section: Unverified / Flagged ----
    if counts["UNVERIFIED"] or counts["FAKE"]:
        doc.add_page_break()
        h = doc.add_heading(
            f"3. Unverified / Flagged References ({counts['UNVERIFIED'] + counts['FAKE']})",
            level=1)
        for r in h.runs:
            r.font.color.rgb = RGBColor.from_string("C00000")
        doc.add_paragraph(
            "⚠ These references did not produce a high-confidence match. They may be:")
        for line in [
            "  • FAKE / HALLUCINATED — the reference does not exist in any database",
            "  • UNVERIFIED — the title or authors are correct but metadata (year, journal) is wrong",
            "  • NOT INDEXED — the reference exists but is not in any of the 81 searched platforms",
            "  • TYPO — the title or DOI has an error preventing a match",
        ]:
            doc.add_paragraph(line)

        flagged = [c for c in classified if c.get("status") in ("UNVERIFIED", "FAKE")]
        for i, c in enumerate(flagged, 1):
            _add_reference_block(doc, c, i, c.get("status", "UNVERIFIED"), _add_badge)

    # ---- Section: Recommendations ----
    doc.add_page_break()
    h = doc.add_heading("4. Recommendations", level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor.from_string("1F4E78")
    recs = []
    if counts["VERIFIED"]:
        recs.append(
            f"• Keep all {counts['VERIFIED']} VERIFIED reference(s) as-is. Their PDFs "
            f"have been downloaded to the 01_PDFs/ subfolder.")
    if counts["LIKELY"]:
        recs.append(
            f"• Manually verify the {counts['LIKELY']} LIKELY reference(s) by clicking "
            f"their DOIs or searching Google Scholar. If confirmed, promote to VERIFIED.")
    if counts["UNVERIFIED"]:
        recs.append(
            f"• {counts['UNVERIFIED']} UNVERIFIED reference(s) need manual review. "
            f"Search Google Scholar with the title + first author + year. If you cannot "
            f"find them, they may be typos or non-existent.")
    if counts["FAKE"]:
        recs.append(
            f"• ⚠ {counts['FAKE']} FAKE reference(s) appear in no database. These are "
            f"very likely hallucinated or fabricated. STRONGLY RECOMMEND: remove from the "
            f"literature review or replace with the actual cited source.")
    for rec in recs:
        doc.add_paragraph(rec)

    # ---- Footer with page numbers ----
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = "Literature Reference Verification Report | Page "
    # Add PAGE field
    run = fp.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    sep_run = fp.add_run(" of ")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText")
    instrText2.set(qn("xml:space"), "preserve")
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    sep_run._r.append(fldChar3)
    sep_run._r.append(instrText2)
    sep_run._r.append(fldChar4)

    # Header
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("Verified by Literature Review Verifier")
    hr.font.size = Pt(9)
    hr.font.italic = True
    hr.font.color.rgb = RGBColor.from_string("888888")

    # Save
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    log.info(f"DOCX report saved: {output_path}")
    return output_path


def _add_reference_block(doc, c: Dict[str, Any], idx: int, status: str,
                         badge_fn) -> None:
    """Add a styled block for one reference to the DOCX."""
    from docx.shared import Pt, RGBColor, Cm
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{idx}. ")
    r.font.bold = True
    r.font.size = Pt(12)

    badge_fn(p, status, STATUS_COLORS[status], STATUS_TEXT_COLORS[status])

    # Reference text
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("Input: ")
    r2.font.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string("666666")
    p2.add_run(c.get("ref", "")[:500]).font.size = Pt(10)

    if c.get("matched_title"):
        p3 = doc.add_paragraph()
        p3.paragraph_format.left_indent = Cm(0.5)
        p3.paragraph_format.space_after = Pt(2)
        r3 = p3.add_run("Matched: ")
        r3.font.bold = True
        r3.font.size = Pt(10)
        r3.font.color.rgb = RGBColor.from_string(STATUS_TEXT_COLORS.get(status, "000000"))
        p3.add_run(f"{c.get('matched_title', '')}").font.size = Pt(11)

    meta_bits = []
    if c.get("matched_authors"):
        meta_bits.append(f"Authors: {c['matched_authors'][:150]}")
    if c.get("matched_year"):
        meta_bits.append(f"Year: {c['matched_year']}")
    if c.get("matched_doi"):
        meta_bits.append(f"DOI: {c['matched_doi']}")
    if c.get("source_platform"):
        meta_bits.append(f"Source: {c['source_platform']}")
    if c.get("source_count"):
        meta_bits.append(f"In {c['source_count']} platform(s)")
    if c.get("score"):
        meta_bits.append(f"AI score: {c['score']:.2f}")
    if meta_bits:
        p4 = doc.add_paragraph()
        p4.paragraph_format.left_indent = Cm(0.5)
        p4.paragraph_format.space_after = Pt(2)
        r4 = p4.add_run(" | ".join(meta_bits))
        r4.font.size = Pt(9)
        r4.font.color.rgb = RGBColor.from_string("555555")

    if c.get("reason"):
        p5 = doc.add_paragraph()
        p5.paragraph_format.left_indent = Cm(0.5)
        p5.paragraph_format.space_after = Pt(2)
        r5 = p5.add_run("Reasoning: ")
        r5.font.bold = True
        r5.font.size = Pt(9)
        r5.font.italic = True
        r5.font.color.rgb = RGBColor.from_string("666666")
        r6 = p5.add_run(c.get("reason", "")[:300])
        r6.font.size = Pt(9)
        r6.font.italic = True
        r6.font.color.rgb = RGBColor.from_string("666666")

    if c.get("download_success"):
        p6 = doc.add_paragraph()
        p6.paragraph_format.left_indent = Cm(0.5)
        p6.paragraph_format.space_after = Pt(8)
        r7 = p6.add_run(f"📥 PDF downloaded: {c.get('download_path', '')}")
        r7.font.size = Pt(9)
        r7.font.color.rgb = RGBColor.from_string("006100")
        r7.font.bold = True
    elif c.get("status") == "VERIFIED" and not c.get("download_success"):
        p6 = doc.add_paragraph()
        p6.paragraph_format.left_indent = Cm(0.5)
        p6.paragraph_format.space_after = Pt(8)
        r7 = p6.add_run("⚠ PDF download attempted but failed (see error log)")
        r7.font.size = Pt(9)
        r7.font.color.rgb = RGBColor.from_string("9C5700")

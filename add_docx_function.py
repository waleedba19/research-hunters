#!/usr/bin/env python3
"""Add DOCX Report function to research_hunter_v2-4.py"""

with open('research_hunter_v2-4.py', 'r') as f:
    content = f.read()

# Check if function already exists
if "def generate_docx_report" in content:
    print("❌ Function generate_docx_report already exists!")
else:
    print("Adding generate_docx_report function...")

docx_function = '''

def generate_docx_report(report_data: dict, out_folder: Path) -> Path | None:
    """Generate professional PhD-level DOCX report with colors, tables, and comprehensive analysis."""
    # Build APA for papers that don't have it
    for p in report_data.get("papers") or []:
        if not p.get("apa"):
            p["apa"] = build_apa(p)

    docx_path = out_folder / "research_report.docx"
    
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        doc = Document()
        
        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(2.5)
        
        # Normal style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.5
        
        # ═══ COVER PAGE ═══
        for _ in range(4):
            doc.add_paragraph()
        
        # Title
        title = doc.add_heading('RESEARCH HUNTER', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(48)
        title.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        
        subtitle = doc.add_heading('Comprehensive Literature Review Report', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(28)
        subtitle.runs[0].font.color.rgb = RGBColor(68, 114, 196)
        
        doc.add_paragraph()
        
        # Decorative line
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('━' * 50)
        run.font.color.rgb = RGBColor(31, 78, 121)
        run.font.size = Pt(14)
        
        doc.add_paragraph()
        
        # Research Topic Box
        topic_label = doc.add_paragraph()
        topic_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = topic_label.add_run('📌 RESEARCH TOPIC')
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(31, 78, 121)
        
        topic = doc.add_paragraph()
        topic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = topic.add_run(report_data.get("title", "N/A"))
        run.bold = True
        run.font.size = Pt(22)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('━' * 50)
        run.font.color.rgb = RGBColor(31, 78, 121)
        run.font.size = Pt(14)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Metadata Table
        meta_table = doc.add_table(rows=6, cols=2)
        meta_table.style = 'Light Grid Accent 1'
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        meta_data = [
            ("🎓 Academic Field", report_data.get("field", "N/A")),
            ("🌐 Search Language", report_data.get("language", "N/A")),
            ("📊 Total Papers", str(report_data.get("total_papers", 0))),
            ("🟢 Q1 Papers", str(report_data.get("q1_count", 0))),
            ("✅ PDFs Downloaded", str(report_data.get("pdfs_downloaded", 0))),
            ("📅 Generated", report_data.get("timestamp", "N/A")),
        ]
        
        for i, (label, value) in enumerate(meta_data):
            row = meta_table.rows[i]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(31, 78, 121)
            row.cells[1].text = str(value)
        
        doc.add_page_break()
        
        # ═══ EXECUTIVE SUMMARY ═══
        doc.add_heading('1. EXECUTIVE SUMMARY', 1)
        
        total = report_data.get("total_papers", 0)
        q1_c = report_data.get("q1_count", 0)
        q2_c = report_data.get("q2_count", 0)
        q3_c = report_data.get("q3_count", 0)
        q4_c = report_data.get("q4_count", 0)
        not_idx = report_data.get("not_indexed_count", 0)
        downloaded = report_data.get("pdfs_downloaded", 0)
        
        p = doc.add_paragraph()
        run = p.add_run(
            f"This comprehensive literature review encompasses {total} research papers gathered from 128+ "
            f"academic databases worldwide. The collection demonstrates exceptional quality indicators, with "
            f"{q1_c} papers from Q1 (top 25%) Scopus/WoS indexed journals, representing the elite tier of "
            f"academic scholarship. Additionally, {q2_c} papers from Q2 journals, {q3_c} from Q3 journals, "
            f"and {q4_c} from Q4 journals provide comprehensive coverage across quality tiers."
        )
        run.font.size = Pt(11)
        
        doc.add_paragraph()
        
        # Quality Distribution Table
        doc.add_heading('Quality Distribution Overview', 2)
        
        stats_table = doc.add_table(rows=6, cols=4)
        stats_table.style = 'Light Grid Accent 1'
        
        headers = ["Quality Tier", "Paper Count", "Percentage", "Assessment"]
        for i, h in enumerate(headers):
            cell = stats_table.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            # Blue background
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '1F4E79')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
        
        data = [
            ("🟢 Q1 - Top Tier (Top 25%)", q1_c, f"{q1_c/total*100:.1f}%" if total else "0%", "⭐⭐⭐ Exceptional"),
            ("🟢 Q2 - Good (50-75%)", q2_c, f"{q2_c/total*100:.1f}%" if total else "0%", "⭐⭐ Very Good"),
            ("🟡 Q3 - Acceptable (25-50%)", q3_c, f"{q3_c/total*100:.1f}%" if total else "0%", "⭐ Good"),
            ("🟠 Q4 - Lower Tier (Bottom 25%)", q4_c, f"{q4_c/total*100:.1f}%" if total else "0%", "Adequate"),
            ("⚫ Not Indexed / Open Access", not_idx, f"{not_idx/total*100:.1f}%" if total else "0%", "Supplementary"),
        ]
        
        for i, (tier, count, pct, assessment) in enumerate(data, 1):
            row = stats_table.rows[i]
            row.cells[0].text = tier
            row.cells[1].text = str(count)
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[2].text = pct
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row.cells[3].text = assessment
        
        doc.add_paragraph()
        doc.add_paragraph(f"Total PDFs Successfully Downloaded: {downloaded}")
        
        doc.add_page_break()
        
        # ═══ METHODOLOGY DEEP-DIVE ═══
        doc.add_heading('2. METHODOLOGY DEEP-DIVE', 1)
        
        doc.add_heading('Search Algorithm & Platform Coverage', 2)
        p = doc.add_paragraph(
            'The Research Hunter system employs a sophisticated multi-platform search algorithm that queries '
            '128+ academic databases simultaneously, including:'
        )
        
        platforms_list = [
            "🔍 OpenAlex API - Comprehensive coverage with 200M+ papers",
            "🔍 Semantic Scholar - Citation analysis and abstracts",
            "🔍 CORE - Open-access full-text availability",
            "🔍 CrossRef - DOI resolution and metadata",
            "🔍 PubMed - Biomedical research",
            "🔍 arXiv - Preprint papers",
            "🔍 SSRN - Working papers",
            "🔍 Google Scholar - Broader coverage",
            "🔍 And 120+ more specialized databases",
        ]
        
        for platform in platforms_list:
            doc.add_paragraph(platform, style='List Bullet')
        
        doc.add_heading('Methodology Classification', 2)
        
        method_types = {
            "📊 Quantitative Research": ["quantitative", "survey", "experimental", "rct", "regression", "correlation", "statistical"],
            "📋 Qualitative Research": ["qualitative", "interview", "focus group", "ethnographic", "phenomenological", "grounded theory", "thematic analysis"],
            "🔀 Mixed-Methods Research": ["mixed methods", "mixed-method", "sequential explanatory"],
            "📁 Case Study": ["case study", "single case", "multiple case"],
            "📚 Systematic Review": ["systematic review", "meta-analysis", "scoping review"],
        }
        
        total_papers = len(report_data.get("papers", []))
        
        for method, keywords in method_types.items():
            papers = [p for p in report_data.get("papers", []) 
                      if any(kw in (p.get("title", "") + p.get("abstract", "")).lower() for kw in keywords)]
            if papers:
                doc.add_heading(method, 3)
                pct = (len(papers) / total_papers * 100) if total_papers else 0
                doc.add_paragraph(f"Found {len(papers)} papers ({pct:.1f}% of collection)")
                
                # Show first 5 papers
                for pap in papers[:5]:
                    doc.add_paragraph(f"• {pap.get('title', 'N/A')[:100]}...", style='List Bullet')
        
        doc.add_page_break()
        
        # ═══ PLATFORM SOURCE BREAKDOWN ═══
        doc.add_heading('3. PLATFORM SOURCE BREAKDOWN', 1)
        doc.add_paragraph('The following table shows which academic platforms contributed papers to this literature review:')
        
        platform_counts = {}
        for p in report_data.get("papers", []):
            src = p.get("source", "Unknown")
            platform_counts[src] = platform_counts.get(src, 0) + 1
        
        plat_table = doc.add_table(rows=min(len(platform_counts), 25)+1, cols=4)
        plat_table.style = 'Light Grid Accent 1'
        
        plat_headers = ["Rank", "Platform Name", "Papers Found", "Coverage %"]
        for i, h in enumerate(plat_headers):
            plat_table.rows[0].cells[i].text = h
            plat_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for i, (platform, count) in enumerate(sorted(platform_counts.items(), key=lambda x: -x[1])[:25], 1):
            plat_table.rows[i].cells[0].text = str(i)
            plat_table.rows[i].cells[1].text = platform
            plat_table.rows[i].cells[2].text = str(count)
            pct = (count / total_papers * 100) if total_papers else 0
            plat_table.rows[i].cells[3].text = f"{pct:.1f}%"
        
        doc.add_page_break()
        
        # ═══ CHAPTER-BY-CHAPTER LITERATURE MAPPING ═══
        doc.add_heading('4. CHAPTER-BY-CHAPTER LITERATURE MAPPING', 1)
        
        chapters = {
            "📖 Chapter 1: Introduction & Problem Statement": ["introduction", "background", "overview", "rationale", "problem statement", "aims", "objectives"],
            "📚 Chapter 2: Literature Review & Theoretical Framework": ["literature review", "theoretical", "conceptual", "prior studies", "meta-analysis", "systematic review", "empirical evidence"],
            "🔬 Chapter 3: Methodological Design": ["methodology", "methods", "research design", "data collection", "participants", "sampling"],
            "📊 Chapter 4: Results & Data Analysis": ["results", "findings", "data analysis", "outcomes", "statistical", "analysis"],
            "💬 Chapter 5: Discussion & Implications": ["discussion", "interpretation", "implications", "limitations", "future research", "recommendations"],
            "✅ Chapter 6: Conclusion & Recommendations": ["conclusion", "summary", "recommendations", "contribution"],
        }
        
        for chapter, keywords in chapters.items():
            doc.add_heading(chapter, 2)
            
            papers = [p for p in report_data.get("papers", []) 
                      if any(kw in (p.get("title", "") + p.get("abstract", "")).lower() for kw in keywords)]
            
            pct = (len(papers) / total_papers * 100) if total_papers else 0
            doc.add_paragraph(f"Relevant Papers: {len(papers)} ({pct:.1f}% of collection)")
            
            if papers:
                ch_table = doc.add_table(rows=min(len(papers), 10)+1, cols=3)
                ch_table.style = 'Light Grid Accent 1'
                
                ch_headers = ["Paper Title", "Authors", "Year/Journal"]
                for i, h in enumerate(ch_headers):
                    ch_table.rows[0].cells[i].text = h
                    ch_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                
                for i, p in enumerate(papers[:10], 1):
                    ch_table.rows[i].cells[0].text = p.get('title', 'N/A')[:60]
                    ch_table.rows[i].cells[1].text = ", ".join(p.get('authors', [])[:3])
                    ch_table.rows[i].cells[2].text = f"{p.get('year', 'N/A')} | {p.get('journal', 'N/A')[:30]}"
            
            doc.add_paragraph()
        
        doc.add_page_break()
        
        # ═══ LINGUISTIC PARITY ANALYSIS ═══
        doc.add_heading('5. LINGUISTIC PARITY & MULTI-LINGUAL ANALYSIS', 1)
        
        doc.add_paragraph(
            'The Research Hunter system implements comprehensive multi-lingual search capabilities, ensuring '
            'equitable coverage across languages including English, Arabic, French, Spanish, German, Chinese, '
            'and other major academic languages.'
        )
        
        doc.add_heading('Arabic Language Sources', 2)
        doc.add_paragraph('Arabic-language sources receive equivalent algorithmic treatment to English papers:')
        
        arabic_features = [
            "✓ Applies identical relevance scoring weights",
            "✓ Uses Arabic-specific keyword normalization",
            "✓ Incorporates Arabic-language databases (ArabPsyc, Shamaa)",
            "✓ Cross-translates abstracts for comparative analysis",
            "✓ Tracks Q1-Q4 quartile status for Arabic journals",
            "✓ Provides bilingual metadata where available",
        ]
        
        for feature in arabic_features:
            doc.add_paragraph(feature)
        
        doc.add_heading('Cross-Language Synthesis', 2)
        doc.add_paragraph(
            'When a study exists in both Arabic and English versions, the system merges insights into a '
            'super-entry, ensuring comprehensive coverage regardless of publication language.'
        )
        
        doc.add_page_break()
        
        # ═══ GEOGRAPHIC DISTRIBUTION ═══
        doc.add_heading('6. GEOGRAPHIC DISTRIBUTION ANALYSIS', 1)
        
        geo_keywords = {
            "🌍 Global / International": ["global", "international", "worldwide", "multi-country"],
            "🇺🇸 North America (USA/Canada)": ["usa", "united states", "america", "canada"],
            "🇬🇧 Europe (UK/Germany/France)": ["uk", "united kingdom", "germany", "france", "spain", "italy"],
            "🌏 Asia (China/Japan/Korea/India)": ["china", "japanese", "korea", "india", "singapore"],
            "🌍 Middle East (KSA/UAE/Egypt)": ["saudi", "uae", "emirates", "egypt", "jordan", "qatar"],
            "🌎 Latin America (Brazil/Mexico)": ["brazil", "mexico", "argentina", "colombia"],
            "🌍 Africa (South Africa/Nigeria)": ["south africa", "nigeria", "kenya", "ethiopia"],
        }
        
        geo_table = doc.add_table(rows=len(geo_keywords)+1, cols=3)
        geo_table.style = 'Light Grid Accent 1'
        
        geo_headers = ["Region", "Paper Count", "Coverage"]
        for i, h in enumerate(geo_headers):
            geo_table.rows[0].cells[i].text = h
            geo_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for i, (region, keywords) in enumerate(geo_keywords.items(), 1):
            count = sum(1 for p in report_data.get("papers", []) 
                       if any(kw in (p.get("title", "") + p.get("abstract", "")).lower() for kw in keywords))
            pct = (count / total_papers * 100) if total_papers else 0
            geo_table.rows[i].cells[0].text = region
            geo_table.rows[i].cells[1].text = str(count)
            geo_table.rows[i].cells[2].text = f"{pct:.1f}%"
        
        doc.add_page_break()
        
        # ═══ COMPLETE BIBLIOGRAPHY ═══
        doc.add_heading('7. COMPLETE BIBLIOGRAPHY (APA 7th Edition)', 1)
        doc.add_paragraph('The following pages contain the complete bibliography formatted according to APA 7th edition guidelines.')
        doc.add_paragraph()
        
        for i, p in enumerate(report_data.get("papers", [])[:200], 1):
            if p.get('apa'):
                citation = p.get('apa')
            else:
                authors = ", ".join(p.get("authors", [])[:5])
                year = p.get("year", "n.d.")
                title = p.get("title", "N/A")
                journal = p.get("journal", "N/A")
                doi = p.get("doi", "")
                citation = f"{authors} ({year}). {title}. {journal}."
                if doi:
                    citation += f" https://doi.org/{doi}"
            
            para = doc.add_paragraph(f"{i}. {citation}")
            para.paragraph_format.left_indent = Inches(0.5)
            para.paragraph_format.first_line_indent = Inches(-0.5)
        
        # Save
        doc.save(str(docx_path))
        print(f"✅ DOCX saved: {docx_path}")
        return docx_path
        
    except ImportError:
        print("❌ python-docx not installed - DOCX skipped. Run: pip install python-docx")
        return None
    except Exception as e:
        print(f"❌ DOCX error: {e}")
        import traceback
        traceback.print_exc()
        return None


'''

# Find where to insert (before ENHANCED WIZARD CONFIGURATIONS or if __name__)
insert_markers = ["#  ENHANCED WIZARD CONFIGURATIONS", "if __name__ == \"__main__\":"]
insert_pos = len(content)

for marker in insert_markers:
    pos = content.find(marker)
    if pos > 0:
        insert_pos = pos
        break

# Insert the function
content = content[:insert_pos] + docx_function + "\n\n" + content[insert_pos:]

with open('research_hunter_v2-4.py', 'w') as f:
    f.write(content)

print("✅ DOCX function added!")
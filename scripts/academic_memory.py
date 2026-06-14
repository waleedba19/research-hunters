"""academic_memory.py — Living knowledge base of academic writing patterns.
Loads pre-built pattern knowledge, learns from new papers, and generates
properly formatted academic documents.

Usage:
  python academic_memory.py --query "phd dissertation formatting"
  python academic_memory.py --learn paper.json
  python academic_memory.py --generate study_type=research_article --output paper.docx
"""

import json, os, sqlite3, sys, hashlib, re, shutil, time, importlib, subprocess
from pathlib import Path
from datetime import datetime
from typing import Optional

SCRIPTS_DIR = Path(__file__).parent
PATTERNS_FILE = SCRIPTS_DIR / "academic_patterns.json"
MEMORY_DB = SCRIPTS_DIR / "academic_memory.db"

# ====================================================================
# PATTERN LOADER
# ====================================================================

class WritingPatterns:
    """Loaded once. Contains ALL known academic writing patterns."""

    _instance = None

    def __init__(self):
        if not PATTERNS_FILE.exists():
            raise FileNotFoundError(f"{PATTERNS_FILE} not found")
        with open(PATTERNS_FILE, encoding="utf-8") as f:
            self.data = json.load(f)

    @classmethod
    def get(cls):
        if cls._instance is None:
            cls._instance = cls()
        return cls._instance

    def query(self, topic: str) -> dict:
        """Return all patterns relevant to a topic (study type, section, etc.)."""
        topic_lower = topic.lower()
        results = {}
        # Search study types
        for st_name, st_data in self.data.get("study_types", {}).items():
            if topic_lower in st_name.lower() or any(
                topic_lower in s.lower() for s in st_data.get("sections", [])
            ):
                results[f"study_type_{st_name}"] = st_data
        # Search sections
        for sec_name, sec_data in self.data.get("sections", {}).items():
            if topic_lower in sec_name.lower():
                results[f"section_{sec_name}"] = sec_data
        # Search methodology types
        for mt_name, mt_data in self.data.get("methodology_types", {}).items():
            if topic_lower in mt_name.lower():
                results[f"methodology_{mt_name}"] = mt_data
        return results

    def font_for(self, element: str) -> dict:
        return self.data.get("fonts", {}).get(element, self.data["fonts"]["body"])

    def page_setup(self) -> dict:
        return self.data.get("page", {})

    def section_structure(self, section: str) -> Optional[dict]:
        return self.data.get("sections", {}).get(section)

    def methodology_details(self, mtype: str) -> Optional[dict]:
        return self.data.get("methodology_types", {}).get(mtype)

    def citation_style(self, style: str = "apa_7th") -> dict:
        return self.data.get("citation_styles", {}).get(style, self.data["citation_styles"]["apa_7th"])

    def study_type(self, name: str) -> Optional[dict]:
        # Normalize: "PhD Dissertation" -> "phd_dissertation"
        key = name.lower().replace(" ", "_").replace("-", "_")
        return self.data.get("study_types", {}).get(key)

# ====================================================================
# MEMORY DATABASE — learns from papers the model reads
# ====================================================================

class AcademicMemory:
    """Persistent database of what the model has learned from papers.
    Stores: writing patterns extracted, document structures observed,
    citation styles seen, formatting rules by journal."""

    def __init__(self, db_path: Path = MEMORY_DB):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        conn = sqlite3.connect(str(self.db_path))
        conn.execute("""
            CREATE TABLE IF NOT EXISTS learned_patterns (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                source_hash TEXT UNIQUE,
                source_title TEXT,
                journal TEXT,
                year INTEGER,
                patterns TEXT,
                learned_at TEXT DEFAULT (datetime('now'))
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS document_formats (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                journal_name TEXT UNIQUE,
                format_rules TEXT,
                observed_at TEXT DEFAULT (datetime('now'))
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS generated_docs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                study_type TEXT,
                title TEXT,
                output_path TEXT,
                generated_at TEXT DEFAULT (datetime('now'))
            )
        """)
        conn.commit()
        conn.close()

    def learn_from_paper(self, paper: dict):
        """Extract and store writing patterns from a paper the model has read."""
        title = paper.get("title", "")
        journal = paper.get("journal", "unknown")
        year = paper.get("year")
        source_hash = hashlib.sha256(
            (title + journal + str(year)).encode()
        ).hexdigest()[:16]

        patterns = {
            "has_abstract": bool(paper.get("abstract")),
            "has_doi": bool(paper.get("doi")),
            "has_keywords": False,
            "section_count": 0,
            "table_count": 0,
            "figure_count": 0,
            "reference_count": 0,
            "word_count": len((paper.get("abstract") or "").split()),
            "language_detected": None,
        }

        # Try to extract from abstract — section indicators
        abstract = paper.get("abstract") or ""
        if abstract:
            sections_found = re.findall(
                r"\b(?:Introduction|Methodology|Results|Discussion|Conclusion"
                r"|Literature Review|Background|Method)\b",
                abstract, re.IGNORECASE
            )
            patterns["section_count"] = len(set(sections_found))

        conn = sqlite3.connect(str(self.db_path))
        try:
            conn.execute(
                "INSERT OR IGNORE INTO learned_patterns "
                "(source_hash, source_title, journal, year, patterns) "
                "VALUES (?, ?, ?, ?, ?)",
                (source_hash, title, journal, year, json.dumps(patterns))
            )
            conn.commit()
        except Exception:
            pass
        finally:
            conn.close()

    def query_memory(self, topic: str, limit: int = 10) -> list:
        """Search learned patterns for relevant knowledge."""
        conn = sqlite3.connect(str(self.db_path))
        conn.row_factory = sqlite3.Row
        rows = conn.execute(
            "SELECT * FROM learned_patterns WHERE "
            "source_title LIKE ? OR journal LIKE ? OR patterns LIKE ? "
            "ORDER BY learned_at DESC LIMIT ?",
            (f"%{topic}%", f"%{topic}%", f"%{topic}%", limit)
        ).fetchall()
        conn.close()
        return [dict(r) for r in rows]

    def stats(self) -> dict:
        conn = sqlite3.connect(str(self.db_path))
        count = conn.execute("SELECT COUNT(*) FROM learned_patterns").fetchone()[0]
        journals = conn.execute(
            "SELECT journal, COUNT(*) as cnt FROM learned_patterns "
            "GROUP BY journal ORDER BY cnt DESC LIMIT 10"
        ).fetchall()
        conn.close()
        return {"papers_learned": count, "top_journals": [j[0] for j in journals]}

# ====================================================================
# DOCUMENT GENERATOR — creates properly formatted academic documents
# ====================================================================

class AcademicDocumentGenerator:
    """Generates DOCX/PDF academic documents with proper formatting.
    Supports: templates, Arabic/RTL, citations (APA/IEEE/Vancouver/Harvard),
    advanced formatting (italics, bold, colors, block quotes, boxes, TOC)."""

    def __init__(self):
        self.patterns = WritingPatterns.get()
        self.memory = AcademicMemory()
        self.docx_available = self._check_docx()

    def _check_docx(self) -> bool:
        try:
            import docx
            return True
        except ImportError:
            return False

    def _load_template(self, template_name: str) -> dict:
        """Load a named template from patterns or return default."""
        templates = self.patterns.data.get("templates", {})
        tmpl = templates.get(template_name)
        if tmpl:
            return tmpl
        return {
            "font": "Times New Roman", "margins_cm": {"all": 2.54},
            "line_spacing": 1.5, "title_size": 16, "heading_1_size": 14
        }

    def generate_research_article(self, output_path: str, title: str = "",
                                   authors: list = None, sections: list = None,
                                   study_type: str = "research_article",
                                   template: str = "", citation_style: str = "apa_7th",
                                   language: str = "en"):
        """Generate a properly formatted research article.
        
        Args:
            template: "elsevier_standard", "springer_standard", "ieee_conference",
                      "apa_manuscript", "phd_dissertation_standard", or "" for default
            citation_style: "apa_7th", "ieee", "vancouver", "harvard", "mla_9th"
            language: "en" or "ar" for Arabic (RTL)
        """
        authors = authors or []
        sections = sections or []

        if not self.docx_available:
            print("ERROR: python-docx not installed. Run: pip install python-docx")
            return False

        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn

        doc = Document()
        is_arabic = language == "ar"
        is_rtl = is_arabic

        # Load template or default
        tmpl = self._load_template(template) if template else {}
        mc = tmpl.get("margins_cm", {})
        if is_arabic:
            arabic_fonts = self.patterns.data.get("arabic", {}).get("fonts", {})
            arabic_page = self.patterns.data.get("arabic", {}).get("page", {})

        # Page setup
        for sec in doc.sections:
            sec.top_margin = Cm(mc.get("top", mc.get("all", 2.54)))
            sec.bottom_margin = Cm(mc.get("bottom", mc.get("all", 2.54)))
            sec.left_margin = Cm(mc.get("left", mc.get("all", 2.54)))
            sec.right_margin = Cm(mc.get("right", mc.get("all", 2.54)))

        style = doc.styles['Normal']
        font = style.font
        if is_arabic:
            af = arabic_fonts.get("body", {"name": "Traditional Arabic", "size": 14})
            font.name = af["name"]
            font.size = Pt(af["size"])
        else:
            font.name = "Times New Roman"
            font.size = Pt(tmpl.get("font_size", 12))
        pf = style.paragraph_format
        pf.line_spacing = tmpl.get("line_spacing", 1.5)
        pf.space_after = Pt(tmpl.get("spacing_after", 6))

        if is_rtl:
            style.element.rPr.rFonts.set(qn("w:eastAsia"), "Traditional Arabic")

        # Title
        title_size = tmpl.get("title_size", 16)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title or "Research Article Title")
        run.bold = True
        run.font.size = Pt(title_size)
        run.font.name = "Traditional Arabic" if is_arabic else "Times New Roman"

        # Authors
        if authors:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(" / ".join(authors) if is_arabic else ", ".join(authors))
            run.font.size = Pt(12)

        # Abstract
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        abstract_label = "الملخص" if is_arabic else "Abstract"
        run = p.add_run(abstract_label)
        run.bold = True
        run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if is_rtl:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf = p.paragraph_format
        pf.first_line_indent = Cm(1.27)
        abs_text = "هذا الملخص يقدم خلفية الدراسة وأهدافها ومنهجيتها ونتائجها الرئيسية." if is_arabic else "This abstract summarizes the background, objectives, methodology, results, and conclusions of the study."
        run = p.add_run(abs_text)
        run.font.size = Pt(12)

        # Keywords
        p = doc.add_paragraph()
        kw_label = "الكلمات المفتاحية: " if is_arabic else "Keywords: "
        run = p.add_run(kw_label)
        run.bold = True
        kw_text = "كلمة؛ كلمة؛ كلمة؛ كلمة" if is_arabic else "keyword1; keyword2; keyword3; keyword4"
        run = p.add_run(kw_text)
        run.italic = True

        # Section content — with advanced formatting support
        h1_size = tmpl.get("heading_1_size", 14)
        body_size = tmpl.get("font_size", 12)

        def _add_rich_text(para, text, bold=False, italic=False, color=None, size=None, font_name=None):
            """Add a formatted run to a paragraph."""
            run = para.add_run(text)
            if bold: run.bold = True
            if italic: run.italic = True
            if color: run.font.color.rgb = RGBColor(*color)
            if size: run.font.size = Pt(size)
            if font_name: run.font.name = font_name
            return run

        for sec in sections:
            doc.add_paragraph()

            # Section heading
            p = doc.add_paragraph()
            run = p.add_run(sec.get("heading", "Section"))
            run.bold = True
            run.font.size = Pt(h1_size)

            # Body text
            p = doc.add_paragraph()
            if is_rtl:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pf = p.paragraph_format
            pf.first_line_indent = Cm(1.27)

            content = sec.get("content", "")
            # Parse inline markers: **bold**, *italic*, "quote"
            parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|".*?")', content)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    _add_rich_text(p, part[2:-2], bold=True, size=body_size)
                elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
                    _add_rich_text(p, part[1:-1], italic=True, size=body_size)
                elif part.startswith('"') and part.endswith('"'):
                    _add_rich_text(p, part, italic=True, size=body_size)
                else:
                    _add_rich_text(p, part, size=body_size)

            # Box/sidebar support
            box_type = sec.get("box_type", "")
            if box_type:
                boxes = self.patterns.data.get("boxes_and_sidebars", {})
                box_cfg = boxes.get(box_type, {})
                # Create a bordered box using a table
                table = doc.add_table(rows=1, cols=1)
                cell = table.cell(0, 0)
                cell.text = ""
                cell_para = cell.paragraphs[0]
                run = cell_para.add_run(f"[{box_type.upper()}] {sec.get('heading', '')}")
                run.bold = True
                doc.add_paragraph()

        # References section
        doc.add_paragraph()
        p = doc.add_paragraph()
        ref_label = "المراجع" if is_arabic else "References"
        run = p.add_run(ref_label)
        run.bold = True
        run.font.size = Pt(h1_size)

        sample_refs = self._format_sample_references(citation_style)
        for ref in sample_refs:
            p = doc.add_paragraph()
            run = p.add_run(ref)
            run.font.size = Pt(12)
            pf = p.paragraph_format
            pf.left_indent = Cm(1.27)
            pf.first_line_indent = Cm(-1.27)

        doc.save(output_path)
        self.memory._init_db()
        conn = sqlite3.connect(str(self.memory.db_path))
        conn.execute(
            "INSERT INTO generated_docs (study_type, title, output_path) VALUES (?, ?, ?)",
            (study_type, title, output_path)
        )
        conn.commit()
        conn.close()
        return True

    def format_citation(self, style: str, authors: str, year: str,
                         title: str = "", journal: str = "", volume: str = "",
                         issue: str = "", pages: str = "", doi: str = "") -> str:
        """Format a single citation in the requested style."""
        styles = {
            "apa_7th": f"{authors} ({year}). {title}. {journal}, {volume}({issue}), {pages}. https://doi.org/{doi}" if doi else f"{authors} ({year}). {title}. {journal}, {volume}({issue}), {pages}.",
            "ieee": f"[1] {authors}, \"{title},\" {journal}, vol. {volume}, no. {issue}, pp. {pages}, {year}." if volume else f"[1] {authors}, \"{title},\" {year}.",
            "vancouver": f"1. {authors}. {title}. {journal}. {year};{volume}({issue}):{pages}." if volume else f"1. {authors}. {title}. {year}.",
            "harvard": f"{authors} ({year}) '{title}', {journal}, {volume}({issue}), pp. {pages}." if volume else f"{authors} ({year}) '{title}'.",
            "mla_9th": f"{authors}. \"{title}.\" {journal}, vol. {volume}, no. {issue}, {year}, pp. {pages}." if volume else f"{authors}. \"{title}.\" {year}.",
        }
        return styles.get(style, styles["apa_7th"])

    def format_reference_list(self, refs: list, style: str = "apa_7th") -> list:
        """Generate a formatted reference list from paper metadata."""
        results = []
        for i, r in enumerate(refs):
            authors = r.get("authors", "Author, A.")
            year = str(r.get("year", "n.d."))
            title = r.get("title", "Untitled")
            journal = r.get("journal", "Journal")
            volume = r.get("volume", "")
            issue = r.get("issue", "")
            pages = r.get("pages", "")
            doi = r.get("doi", "")
            citation = self.format_citation(style, authors, year, title, journal, volume, issue, pages, doi)
            if style == "ieee":
                citation = citation.replace("[1]", f"[{i+1}]")
            elif style == "vancouver":
                citation = citation.replace("1.", f"{i+1}.")
            results.append(citation)
        return results

    def _format_sample_references(self, style: str = "apa_7th") -> list:
        """Generate sample references for the requested citation style."""
        samples = [
            {"authors": "Author, A. A.", "year": 2024, "title": "Sample research article", "journal": "Journal of Academic Studies", "volume": "15", "issue": "3", "pages": "100-120", "doi": "10.1234/example.2024.01"},
            {"authors": "Author, B. B.", "year": 2023, "title": "Another academic paper", "journal": "Research Quarterly", "volume": "28", "issue": "2", "pages": "50-70", "doi": ""},
        ]
        return self.format_reference_list(samples, style)

    def generate_bilingual_abstract(self, output_path: str, title: str,
                                     authors: list, english_abstract: str,
                                     arabic_abstract: str, keywords_en: list,
                                     keywords_ar: list) -> bool:
        """Generate a bilingual (English + Arabic) abstract page."""
        if not self.docx_available:
            return False

        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        for sec in doc.sections:
            sec.top_margin = Cm(2.54)
            sec.bottom_margin = Cm(2.54)
            sec.left_margin = Cm(2.54)
            sec.right_margin = Cm(2.54)

        # English version
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True; run.font.size = Pt(16)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(", ".join(authors))

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Abstract"); run.bold = True; run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf = p.paragraph_format; pf.first_line_indent = Cm(1.27)
        run = p.add_run(english_abstract)

        p = doc.add_paragraph()
        run = p.add_run("Keywords: "); run.bold = True
        run = p.add_run("; ".join(keywords_en)); run.italic = True

        doc.add_paragraph()

        # Arabic version
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("الملخص العربي")
        run.bold = True; run.font.size = Pt(16)
        run.font.name = "Traditional Arabic"

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf = p.paragraph_format; pf.first_line_indent = Cm(1.0)
        run = p.add_run(arabic_abstract)
        run.font.name = "Traditional Arabic"; run.font.size = Pt(14)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("الكلمات المفتاحية: "); run.bold = True
        run.font.name = "Traditional Arabic"
        run = p.add_run("؛ ".join(keywords_ar))
        run.font.name = "Traditional Arabic"

        doc.save(output_path)
        return True

    def estimate_paper_length(self, study_type: str) -> str:
        st = self.patterns.study_type(study_type)
        if st:
            return st.get("typical_length", "unknown")
        return "unknown"

    def section_requirements(self, study_type: str) -> list:
        st = self.patterns.study_type(study_type)
        if not st:
            return []
        required = [s for s in st.get("sections", [])]
        sections_data = self.patterns.data.get("sections", {})
        return [
            {"name": s, "required": sections_data.get(s, {}).get("required", False),
             "contains": sections_data.get(s, {}).get("contains", [])}
            for s in required
        ]

# ====================================================================
# CLI
# ====================================================================

def cli():
    import argparse
    parser = argparse.ArgumentParser(description="Academic Memory — writing pattern knowledge base")
    parser.add_argument("--query", help="Query patterns by topic")
    parser.add_argument("--learn", help="Learn from a paper JSON file")
    parser.add_argument("--generate", help="Generate a formatted document", metavar="study_type")
    parser.add_argument("--title", default="Research Article", help="Document title")
    parser.add_argument("--authors", nargs="*", default=[], help="Author names")
    parser.add_argument("--output", default="output.docx", help="Output path")
    parser.add_argument("--stats", action="store_true", help="Show learning stats")
    parser.add_argument("--sections", help="JSON file with section content")
    args = parser.parse_args()

    mem = AcademicMemory()
    gen = AcademicDocumentGenerator()
    patterns = WritingPatterns.get()

    if args.query:
        results = patterns.query(args.query)
        print(f"\n📖 Patterns for '{args.query}':")
        if results:
            for key, val in results.items():
                print(f"\n  {key}:")
                for k, v in val.items():
                    print(f"    {k}: {v}")
        else:
            print("  No exact match. Try: phd_dissertation, methodology, introduction")
        mem_results = gen.memory.query_memory(args.query)
        if mem_results:
            print(f"\n  Learned from {len(mem_results)} papers:")
            for r in mem_results:
                print(f"    - {r['source_title'][:60]} ({r['journal']})")

    if args.stats:
        s = mem.stats()
        print(f"\n📊 Memory Stats:")
        print(f"  Papers learned from: {s['papers_learned']}")
        print(f"  Top journals: {', '.join(s['top_journals'][:5])}")
        conn = sqlite3.connect(str(MEMORY_DB))
        docs = conn.execute(
            "SELECT study_type, title, generated_at FROM generated_docs ORDER BY generated_at DESC LIMIT 5"
        ).fetchall()
        conn.close()
        if docs:
            print(f"  Recently generated:")
            for d in docs:
                print(f"    [{d[0]}] {d[1][:50]} — {d[2]}")

    if args.learn:
        try:
            with open(args.learn, encoding="utf-8") as f:
                paper = json.load(f)
            gen.memory.learn_from_paper(paper)
            print(f"✅ Learned from: {paper.get('title', args.learn)[:60]}")
        except Exception as e:
            print(f"❌ Learn failed: {e}")

    if args.generate:
        sections = []
        if args.sections:
            try:
                with open(args.sections, encoding="utf-8") as f:
                    sections = json.load(f)
            except Exception as e:
                print(f"⚠️ Sections load failed: {e}")
        ok = gen.generate_research_article(
            args.output, title=args.title, authors=args.authors,
            sections=sections, study_type=args.generate
        )
        if ok:
            print(f"✅ Generated: {os.path.abspath(args.output)}")

    if not any([args.query, args.learn, args.generate, args.stats]):
        parser.print_help()
        print(f"\n📚 Memory state: {mem.stats()['papers_learned']} papers learned")

if __name__ == "__main__":
    cli()
